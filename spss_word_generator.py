"""
SPSS Word Generator - Academic Standard for Algerian Theses
Version: 2.5 - Enhanced for Algerian Academic Standards
Date: December 2024

Features:
- Complete methodological information
- Descriptive statistics tables
- Post-hoc tests for ANOVA
- Mathematical equations for regression
- Extended academic interpretations
- Writing guide for thesis chapters
- RTL support for Arabic text
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


class SPSSWordGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """Setup document margins and defaults"""
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def _add_title(self, text, level=1):
        """Add formatted title with RTL support"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.right_to_left = True  # RTL fix
        
        run = title.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        title.paragraph_format.space_after = Pt(12)
        return title
    
    def _add_section_header(self, text):
        """Add section header with RTL support"""
        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.right_to_left = True  # RTL fix
        
        run = header.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        header.paragraph_format.space_before = Pt(12)
        header.paragraph_format.space_after = Pt(6)
        return header
    
    def _add_paragraph(self, text, align='right', bold=False):
        """Add formatted paragraph with RTL support"""
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if align == 'right' else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.right_to_left = True  # RTL fix
        
        run = para.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = bold
        
        return para
    
    def _create_table(self, rows, cols, headers=None):
        """Create formatted table"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'
        
        if headers:
            for i, header_text in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)
                        run.font.name = 'Times New Roman'
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return table
    
    def _fill_table_cell(self, cell, text, align='center', bold=False):
        """Fill table cell with formatted text"""
        cell.text = str(text)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.bold = bold
    
    def generate_anova(self, results):
        """Generate One-Way ANOVA report - Enhanced for Algerian Standards"""
        self._add_title("تحليل التباين الأحادي\nOne-Way ANOVA")
        self.doc.add_paragraph()
        
        if 'error' in results:
            self._add_paragraph(f"❌ خطأ: {results['error']}")
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 معلومات التحليل:")
        self._add_paragraph(f"• الاختبار: تحليل التباين الأحادي (One-Way ANOVA)")
        if 'إحصاءات_المجموعات' in results:
            self._add_paragraph(f"• عدد المجموعات: {len(results['إحصاءات_المجموعات'])}")
        self._add_paragraph(f"• العدد الكلي: N = {results.get('N', 'غير محدد')}")
        self._add_paragraph(f"• مستوى الدلالة: α = 0.05")
        self.doc.add_paragraph()
        
        # الإحصاءات الوصفية للمجموعات
        self._add_section_header("📊 أولاً: الإحصاءات الوصفية للمجموعات")
        self._add_paragraph("يعرض الجدول التالي الإحصاءات الوصفية لكل مجموعة من مجموعات المتغير المستقل، "
                           "مما يساعد في فهم توزيع البيانات والفروق الظاهرية بين المجموعات قبل التحليل الإحصائي.")
        self.doc.add_paragraph()
        
        if 'إحصاءات_المجموعات' in results:
            groups = results['إحصاءات_المجموعات']
            table = self._create_table(rows=len(groups) + 1, cols=4, headers=['المجموعة', 'N', 'Mean', 'Std. Deviation'])
            for i, (group_name, stats) in enumerate(groups.items(), start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], str(group_name), align='right', bold=True)
                self._fill_table_cell(cells[1], stats.get('العدد', '-'))
                self._fill_table_cell(cells[2], f"{stats.get('المتوسط', 0):.2f}")
                self._fill_table_cell(cells[3], f"{stats.get('الانحراف_المعياري', 0):.2f}")
            self.doc.add_paragraph()
        
        # جدول تحليل التباين
        self._add_section_header("📈 ثانياً: جدول تحليل التباين ANOVA")
        self._add_paragraph("يوضح الجدول التالي نتائج تحليل التباين الأحادي، حيث يتم مقارنة التباين بين المجموعات "
                           "بالتباين داخل المجموعات للكشف عن الفروق ذات الدلالة الإحصائية.")
        self.doc.add_paragraph()
        
        table = self._create_table(rows=4, cols=6, headers=['مصدر التباين', 'Sum of Squares', 'df', 'Mean Square', 'F', 'Sig.'])
        
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], 'بين المجموعات', align='right')
        self._fill_table_cell(cells[1], f"{results['بين_المجموعات']['مجموع_المربعات']:.3f}")
        self._fill_table_cell(cells[2], results['بين_المجموعات']['درجات_الحرية'])
        self._fill_table_cell(cells[3], f"{results['بين_المجموعات']['متوسط_المربعات']:.3f}")
        self._fill_table_cell(cells[4], f"{results['F']:.3f}")
        self._fill_table_cell(cells[5], f"{results['p']:.4f}")
        
        cells = table.rows[2].cells
        self._fill_table_cell(cells[0], 'داخل المجموعات', align='right')
        self._fill_table_cell(cells[1], f"{results['داخل_المجموعات']['مجموع_المربعات']:.3f}")
        self._fill_table_cell(cells[2], results['داخل_المجموعات']['درجات_الحرية'])
        self._fill_table_cell(cells[3], f"{results['داخل_المجموعات']['متوسط_المربعات']:.3f}")
        self._fill_table_cell(cells[4], '-')
        self._fill_table_cell(cells[5], '-')
        
        cells = table.rows[3].cells
        self._fill_table_cell(cells[0], 'المجموع', align='right')
        self._fill_table_cell(cells[1], f"{results['الكلي']['مجموع_المربعات']:.3f}")
        self._fill_table_cell(cells[2], results['الكلي']['درجات_الحرية'])
        self._fill_table_cell(cells[3], '-')
        self._fill_table_cell(cells[4], '-')
        self._fill_table_cell(cells[5], '-')
        
        self.doc.add_paragraph()
        
        # Post-hoc Tests (عند وجود دلالة)
        if 'post_hoc' in results and results.get('دال', False):
            self._add_section_header("📊 ثالثاً: المقارنات البعدية (Post-hoc Tests)")
            self._add_paragraph(
                f"نظراً لوجود فروق دالة إحصائياً في اختبار ANOVA، تم إجراء المقارنات البعدية "
                f"باستخدام طريقة {results['post_hoc']['method']} لتحديد أي المجموعات تختلف بشكل دال عن الأخرى. "
                f"تُستخدم هذه الطريقة لضبط مستوى الدلالة عند إجراء مقارنات متعددة، مما يقلل من احتمالية الخطأ من النوع الأول."
            )
            self.doc.add_paragraph()
            
            comparisons = results['post_hoc']['comparisons']
            table = self._create_table(
                rows=len(comparisons) + 1,
                cols=4,
                headers=['المجموعة (I)', 'المجموعة (J)', 'فرق المتوسطات (I-J)', 'Sig.']
            )
            
            for i, comp in enumerate(comparisons, start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], comp['group1'], align='right', bold=True)
                self._fill_table_cell(cells[1], comp['group2'], align='right', bold=True)
                self._fill_table_cell(cells[2], f"{comp['mean_diff']:.3f}")
                sig_text = f"{comp['p']:.4f}"
                if comp['دال']:
                    sig_text += "*"
                self._fill_table_cell(cells[3], sig_text)
            
            self.doc.add_paragraph()
            
            # تفسير المقارنات الدالة
            dalah_comps = [c for c in comparisons if c['دال']]
            if dalah_comps:
                interp = "من خلال جدول المقارنات البعدية أعلاه، يتضح وجود فروق دالة إحصائياً بين المجموعات التالية:\n\n"
                for comp in dalah_comps:
                    direction = "أعلى" if comp['mean_diff'] > 0 else "أقل"
                    interp += f"• الفرق بين مجموعة ({comp['group1']}) ومجموعة ({comp['group2']}): حيث كان متوسط مجموعة {comp['group1']} {direction} بفارق قدره ({abs(comp['mean_diff']):.2f}) درجة، وهو فرق دال إحصائياً عند مستوى (p = {comp['p']:.4f}).\n\n"
                self._add_paragraph(interp)
            else:
                self._add_paragraph(
                    "بالرغم من وجود فروق دالة إحصائياً في اختبار ANOVA الأساسي، إلا أن المقارنات البعدية "
                    "لم تُظهر فروقاً دالة بين أي مجموعتين عند تطبيق التصحيح الإحصائي للمقارنات المتعددة. "
                    "وهذا يُعزى إلى أن التصحيح الإحصائي (مثل Bonferroni) يرفع معيار الدلالة المطلوب، "
                    "مما قد يؤدي إلى عدم ظهور فروق دالة بين أزواج المجموعات الفردية رغم وجود فروق عامة."
                )
        
        self.doc.add_paragraph()
        
        # التفسير الأكاديمي المطول
        section_number = "رابعاً" if 'post_hoc' in results and results.get('دال') else "ثالثاً"
        self._add_section_header(f"📖 {section_number}: التفسير الأكاديمي المفصل")
        
        if results['دال']:
            df_b = results['بين_المجموعات']['درجات_الحرية']
            df_w = results['داخل_المجموعات']['درجات_الحرية']
            
            interp = (
                f"أظهرت نتائج تحليل التباين الأحادي (One-Way ANOVA) وجود فروق ذات دلالة إحصائية بين المجموعات "
                f"المدروسة عند مستوى دلالة {results['مستوى_الدلالة']}, حيث بلغت قيمة F المحسوبة ({results['F']:.3f}) "
                f"بدرجات حرية ({df_b}, {df_w}), وبقيمة احتمالية p = {results['p']:.4f}. "
                f"وبما أن قيمة p أقل من مستوى الدلالة المعتمد (0.05)، فإننا نرفض الفرضية الصفرية ونقبل الفرضية البديلة، "
                f"مما يعني وجود فروق جوهرية بين متوسطات المجموعات.\n\n"
                f"كما بلغ حجم الأثر (Eta Squared = {results['eta_squared']:.3f}) وهو يُصنف على أنه {results['حجم_الأثر']}، "
                f"مما يشير إلى أن المتغير المستقل يفسر ما نسبته {results['eta_squared']*100:.1f}% من التباين الكلي "
                f"في المتغير التابع. وهذا يدل على وجود أثر عملي ملموس للمتغير المستقل على المتغير التابع، "
                f"وليس مجرد دلالة إحصائية فقط.\n\n"
                f"من الناحية العملية، تشير هذه النتائج إلى أن الاختلافات بين المجموعات ليست عشوائية، "
                f"وإنما تعكس تأثيراً حقيقياً للمتغير المستقل. ويمكن الاعتماد على هذه النتائج في اتخاذ القرارات "
                f"أو بناء التوصيات المتعلقة بموضوع الدراسة."
            )
        else:
            df_b = results['بين_المجموعات']['درجات_الحرية']
            df_w = results['داخل_المجموعات']['درجات_الحرية']
            
            interp = (
                f"أظهرت نتائج تحليل التباين الأحادي (One-Way ANOVA) عدم وجود فروق ذات دلالة إحصائية "
                f"بين المجموعات المدروسة عند مستوى دلالة 0.05, حيث بلغت قيمة F المحسوبة ({results['F']:.3f}) "
                f"بدرجات حرية ({df_b}, {df_w}), وبقيمة احتمالية p = {results['p']:.4f}. "
                f"وبما أن قيمة p أكبر من مستوى الدلالة المعتمد (0.05)، فإننا نقبل الفرضية الصفرية، "
                f"مما يعني عدم وجود فروق جوهرية بين متوسطات المجموعات.\n\n"
                f"وهذا يشير إلى أن المتغير المستقل لم يُظهر تأثيراً دالاً إحصائياً على المتغير التابع في هذه العينة. "
                f"ومع ذلك، يجب الأخذ بعين الاعتبار أن عدم وجود دلالة إحصائية لا يعني بالضرورة عدم وجود فروق فعلية، "
                f"بل قد يعود ذلك إلى محدودية حجم العينة، أو وجود تداخل كبير بين المجموعات، أو تأثير عوامل أخرى "
                f"لم تُضبط في الدراسة.\n\n"
                f"من الناحية العملية، تشير هذه النتائج إلى تشابه المجموعات المدروسة في المتغير التابع، "
                f"مما قد يدعو إلى إعادة النظر في الفرضيات أو تصميم الدراسة، أو البحث عن متغيرات أخرى قد تفسر "
                f"التباين في المتغير التابع بشكل أفضل."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        next_section = "خامساً" if 'post_hoc' in results and results.get('دال') else "رابعاً"
        self._add_section_header(f"📝 {next_section}: كيفية الكتابة في المذكرة")
        
        self._add_paragraph("• في فصل الإجراءات المنهجية:", bold=True)
        self._add_paragraph(
            f'"تم استخدام اختبار تحليل التباين الأحادي (One-Way ANOVA) للكشف عن الفروق بين المجموعات، '
            f'حيث بلغت العينة الكلية N = {results.get("N", "X")}. وقد تم اعتماد مستوى دلالة α = 0.05 '
            f'كمعيار للحكم على الدلالة الإحصائية."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        if results['دال']:
            self._add_paragraph(
                '"أظهرت نتائج تحليل التباين الأحادي وجود فروق دالة إحصائياً بين المجموعات '
                '(F = X.XX, p < 0.05), مما يدل على تأثير [المتغير المستقل] على [المتغير التابع]. '
                'وقد بلغ حجم الأثر (η² = X.XX) مما يشير إلى تأثير [ضعيف/متوسط/كبير]."'
            )
        else:
            self._add_paragraph(
                '"أظهرت نتائج تحليل التباين الأحادي عدم وجود فروق دالة إحصائياً بين المجموعات '
                '(F = X.XX, p > 0.05), مما يشير إلى تشابه المجموعات في [المتغير التابع]."'
            )
        
        return self.doc
    
    def generate_correlation(self, results):
        """Generate Correlation Analysis report"""
        self._add_title("تحليل الارتباط\nCorrelation Analysis")
        self.doc.add_paragraph()
        
        if 'error' in results:
            self._add_paragraph(f"❌ خطأ: {results['error']}")
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 معلومات التحليل:")
        method_ar = "بيرسون" if results.get('method') == 'pearson' else "سبيرمان"
        method_en = "Pearson" if results.get('method') == 'pearson' else "Spearman"
        self._add_paragraph(f"• الاختبار: معامل ارتباط {method_ar} ({method_en} Correlation)")
        self._add_paragraph(f"• العدد الكلي: N = {results.get('N', 'غير محدد')}")
        self._add_paragraph(f"• مستوى الدلالة: α = 0.05")
        self.doc.add_paragraph()
        
        # الإحصاءات الوصفية
        self._add_section_header("📊 أولاً: الإحصاءات الوصفية للمتغيرات")
        self._add_paragraph("يعرض الجدول التالي الإحصاءات الوصفية للمتغيرات المدروسة في تحليل الارتباط، "
                           "مما يساعد في فهم خصائص توزيع كل متغير قبل دراسة العلاقات بينها.")
        self.doc.add_paragraph()
        
        if 'إحصاءات_وصفية' in results:
            descriptives = results['إحصاءات_وصفية']
            table = self._create_table(rows=len(descriptives) + 1, cols=4, headers=['المتغير', 'N', 'Mean', 'Std. Deviation'])
            for i, (var_name, stats) in enumerate(descriptives.items(), start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], str(var_name), align='right', bold=True)
                self._fill_table_cell(cells[1], stats.get('N', '-'))
                self._fill_table_cell(cells[2], f"{stats.get('Mean', 0):.2f}")
                self._fill_table_cell(cells[3], f"{stats.get('SD', 0):.2f}")
            self.doc.add_paragraph()
        
        # مصفوفة الارتباط
        self._add_section_header("📈 ثانياً: مصفوفة الارتباط")
        self._add_paragraph(
            "يعرض الجدول التالي معاملات الارتباط بين جميع أزواج المتغيرات، حيث تشير النجوم إلى مستوى "
            "الدلالة الإحصائية (* p < 0.05, ** p < 0.01, *** p < 0.001). وتتراوح قيم معامل الارتباط "
            "بين -1 (ارتباط سالب تام) و +1 (ارتباط موجب تام)، حيث تشير القيمة 0 إلى عدم وجود ارتباط خطي."
        )
        self.doc.add_paragraph()
        
        if 'مصفوفة_الارتباط' in results:
            matrix = results['مصفوفة_الارتباط']
            variables = list(matrix.keys())
            table = self._create_table(rows=len(variables) + 1, cols=len(variables) + 1, headers=[''] + variables)
            
            for i, var1 in enumerate(variables, start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], var1, align='right', bold=True)
                for j, var2 in enumerate(variables, start=1):
                    r_value = matrix[var1][var2]['r']
                    p_value = matrix[var1][var2]['p']
                    if p_value < 0.001:
                        sig_text = f"{r_value:.3f}***"
                    elif p_value < 0.01:
                        sig_text = f"{r_value:.3f}**"
                    elif p_value < 0.05:
                        sig_text = f"{r_value:.3f}*"
                    else:
                        sig_text = f"{r_value:.3f}"
                    self._fill_table_cell(cells[j], sig_text)
            
            self.doc.add_paragraph()
            
            # Note about N
            note = self.doc.add_paragraph()
            note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            note.paragraph_format.right_to_left = True
            run = note.add_run(f"Note: N = {results.get('N', 'X')} for all correlations.")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.italic = True
            self.doc.add_paragraph()
        
        # التفسير الأكاديمي المطول
        self._add_section_header("📖 ثالثاً: التفسير الأكاديمي المفصل")
        
        if 'نتائج_دالة' in results and results['نتائج_دالة']:
            interp = (
                "أظهرت نتائج تحليل الارتباط باستخدام معامل ارتباط " + method_ar + 
                " وجود علاقات ذات دلالة إحصائية بين بعض المتغيرات المدروسة. وفيما يلي تفصيل لأهم "
                "العلاقات الارتباطية الدالة:\n\n"
            )
            
            for result in results['نتائج_دالة']:
                direction = "موجبة (طردية)" if result['r'] > 0 else "سالبة (عكسية)"
                strength = result.get('قوة', 'متوسطة')
                
                interp += (
                    f"• العلاقة بين {result['var1']} و {result['var2']}: أظهرت النتائج وجود علاقة ارتباطية "
                    f"{direction} وذات قوة {strength} بين المتغيرين، حيث بلغ معامل الارتباط "
                    f"(r = {result['r']:.3f}) وهو دال إحصائياً عند مستوى (p = {result['p']:.4f}). "
                    f"وهذا يعني أن {'الزيادة' if result['r'] > 0 else 'النقصان'} في أحد المتغيرين "
                    f"{'يرتبط بزيادة' if result['r'] > 0 else 'يرتبط بنقصان'} في المتغير الآخر بدرجة {strength}.\n\n"
                )
            
            interp += (
                "\n\nمن الناحية العملية، تشير هذه النتائج إلى وجود علاقات معنوية بين المتغيرات، "
                "مما يمكن الباحثين من فهم طبيعة العلاقات بين المتغيرات المدروسة. ومع ذلك، يجب التنبيه "
                "إلى أن الارتباط لا يعني بالضرورة وجود علاقة سببية، بل يشير فقط إلى وجود علاقة خطية "
                "بين المتغيرات، والتي قد تكون ناتجة عن تأثير متغيرات أخرى غير مدروسة."
            )
        else:
            interp = (
                "أظهرت نتائج تحليل الارتباط باستخدام معامل ارتباط " + method_ar + 
                " عدم وجود علاقات ذات دلالة إحصائية بين المتغيرات المدروسة عند مستوى دلالة 0.05. "
                "وهذا يشير إلى أن المتغيرات المدروسة قد تكون مستقلة عن بعضها البعض، أو أن العلاقة بينها "
                "ضعيفة جداً بحيث لا يمكن اكتشافها بحجم العينة الحالي.\n\n"
                "من الناحية العملية، قد يشير هذا إلى الحاجة لإعادة النظر في اختيار المتغيرات، أو زيادة "
                "حجم العينة، أو البحث عن علاقات غير خطية قد تكون موجودة بين المتغيرات."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        self._add_section_header("📝 رابعاً: كيفية الكتابة في المذكرة")
        
        self._add_paragraph("• في فصل الإجراءات المنهجية:", bold=True)
        self._add_paragraph(
            f'"تم استخدام معامل ارتباط {method_ar} ({method_en}) لقياس قوة واتجاه العلاقة بين المتغيرات، '
            f'حيث بلغت العينة N = {results.get("N", "X")}. وقد تم اعتماد مستوى دلالة α = 0.05 '
            f'كمعيار للحكم على الدلالة الإحصائية للارتباطات."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        self._add_paragraph(
            '"أظهرت نتائج تحليل الارتباط وجود علاقة [موجبة/سالبة] [ضعيفة/متوسطة/قوية] ذات دلالة إحصائية '
            'بين [المتغير الأول] و[المتغير الثاني] (r = X.XX, p < 0.05)، مما يشير إلى أن [تفسير العلاقة]."'
        )
        
        return self.doc
    
    def generate_chisquare(self, results):
        """Generate Chi-Square Test report"""
        self._add_title("اختبار مربع كاي\nChi-Square Test")
        self.doc.add_paragraph()
        
        if 'error' in results:
            self._add_paragraph(f"❌ خطأ: {results['error']}")
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 معلومات التحليل:")
        self._add_paragraph(f"• الاختبار: اختبار مربع كاي للاستقلالية (Chi-Square Test of Independence)")
        self._add_paragraph(f"• المتغير الأول: {results.get('var1', 'غير محدد')}")
        self._add_paragraph(f"• المتغير الثاني: {results.get('var2', 'غير محدد')}")
        self._add_paragraph(f"• العدد الكلي: N = {results.get('N', 'غير محدد')}")
        self._add_paragraph(f"• مستوى الدلالة: α = 0.05")
        self.doc.add_paragraph()
        
        # جدول التوافق
        self._add_section_header("📊 أولاً: جدول التوافق (Crosstabulation)")
        self._add_paragraph(
            "يعرض الجدول التالي التوزيع التكراري المشترك للحالات حسب فئات المتغيرين المدروسين، "
            "مما يساعد في فهم كيفية توزع الحالات عبر مختلف التقاطعات بين فئات المتغيرين. "
            "وتُستخدم هذه البيانات لحساب قيمة مربع كاي واختبار الاستقلالية."
        )
        self.doc.add_paragraph()
        
        if 'جدول_التوافق' in results:
            crosstab = results['جدول_التوافق']
            row_categories = list(crosstab.keys())
            col_categories = list(crosstab[row_categories[0]].keys())
            
            table = self._create_table(
                rows=len(row_categories) + 2,
                cols=len(col_categories) + 2,
                headers=[''] + col_categories + ['المجموع']
            )
            
            col_totals = {col: 0 for col in col_categories}
            grand_total = 0
            
            for i, row_cat in enumerate(row_categories, start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], str(row_cat), align='right', bold=True)
                row_total = 0
                for j, col_cat in enumerate(col_categories, start=1):
                    count = crosstab[row_cat][col_cat]
                    self._fill_table_cell(cells[j], str(count))
                    row_total += count
                    col_totals[col_cat] += count
                self._fill_table_cell(cells[-1], str(row_total), bold=True)
                grand_total += row_total
            
            last_row_cells = table.rows[-1].cells
            self._fill_table_cell(last_row_cells[0], 'المجموع', align='right', bold=True)
            for j, col_cat in enumerate(col_categories, start=1):
                self._fill_table_cell(last_row_cells[j], str(col_totals[col_cat]), bold=True)
            self._fill_table_cell(last_row_cells[-1], str(grand_total), bold=True)
            
            self.doc.add_paragraph()
        
        # نتائج Chi-Square
        self._add_section_header("📈 ثانياً: نتائج اختبار مربع كاي")
        self._add_paragraph(
            "يعرض الجدول التالي نتائج اختبار مربع كاي للاستقلالية، والذي يختبر ما إذا كان هناك "
            "علاقة دالة إحصائياً بين المتغيرين الاسميين أم أن المتغيرين مستقلان عن بعضهما البعض."
        )
        self.doc.add_paragraph()
        
        table = self._create_table(
            rows=2,
            cols=4,
            headers=['Chi-Square (χ²)', 'df', 'Asymp. Sig.', "Cramér's V"]
        )
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], f"{results['chi_square']:.3f}")
        self._fill_table_cell(cells[1], results['df'])
        self._fill_table_cell(cells[2], f"{results['p']:.4f}")
        if 'cramers_v' in results:
            self._fill_table_cell(cells[3], f"{results['cramers_v']:.3f}")
        else:
            self._fill_table_cell(cells[3], '-')
        
        self.doc.add_paragraph()
        
        # التفسير الأكاديمي المطول
        self._add_section_header("📖 ثالثاً: التفسير الأكاديمي المفصل")
        
        if results.get('دال'):
            interp = (
                f"أظهرت نتائج اختبار مربع كاي للاستقلالية وجود علاقة ذات دلالة إحصائية بين المتغيرين "
                f"({results.get('var1', 'المتغير الأول')}) و ({results.get('var2', 'المتغير الثاني')}) "
                f"عند مستوى دلالة {results.get('مستوى_الدلالة', '0.05')}. حيث بلغت قيمة مربع كاي المحسوبة "
                f"(χ² = {results['chi_square']:.3f}) بدرجات حرية (df = {results['df']}), "
                f"وبقيمة احتمالية (p = {results['p']:.4f}).\n\n"
            )
            
            if 'cramers_v' in results:
                strength = results.get('قوة_العلاقة', 'متوسطة')
                interp += (
                    f"كما بلغت قيمة معامل كرامر (Cramér's V = {results['cramers_v']:.3f}), وهو مقياس "
                    f"لقوة العلاقة بين المتغيرين الاسميين، ويشير هذا المعامل إلى وجود علاقة {strength} "
                    f"بين المتغيرين. ويتراوح هذا المعامل بين 0 (عدم وجود علاقة) و 1 (علاقة تامة).\n\n"
                )
            
            interp += (
                "من الناحية العملية، تشير هذه النتائج إلى أن توزيع الحالات عبر فئات المتغير الأول "
                "يختلف باختلاف فئات المتغير الثاني، وليس مجرد توزيع عشوائي. وبالتالي، فإن معرفة فئة "
                "أحد المتغيرين تساعد في التنبؤ بفئة المتغير الآخر. وهذا يعني وجود ارتباط أو علاقة "
                "تبعية بين المتغيرين، مما قد يكون له أهمية نظرية أو تطبيقية حسب موضوع الدراسة."
            )
        else:
            interp = (
                f"أظهرت نتائج اختبار مربع كاي للاستقلالية عدم وجود علاقة ذات دلالة إحصائية بين المتغيرين "
                f"({results.get('var1', 'المتغير الأول')}) و ({results.get('var2', 'المتغير الثاني')}) "
                f"عند مستوى دلالة 0.05. حيث بلغت قيمة مربع كاي المحسوبة (χ² = {results['chi_square']:.3f}) "
                f"بدرجات حرية (df = {results['df']}), وبقيمة احتمالية (p = {results['p']:.4f}), "
                f"وهي قيمة أكبر من مستوى الدلالة المعتمد (0.05).\n\n"
                "من الناحية العملية، تشير هذه النتائج إلى أن المتغيرين مستقلان عن بعضهما البعض، "
                "أي أن توزيع الحالات عبر فئات المتغير الأول لا يتأثر بفئات المتغير الثاني. "
                "وبالتالي، فإن معرفة فئة أحد المتغيرين لا تساعد في التنبؤ بفئة المتغير الآخر. "
                "وهذا قد يشير إلى أن المتغيرين لا يرتبطان ببعضهما في هذه العينة، أو أن حجم العينة "
                "غير كافٍ للكشف عن علاقة ضعيفة قد تكون موجودة."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        self._add_section_header("📝 رابعاً: كيفية الكتابة في المذكرة")
        
        self._add_paragraph("• في فصل الإجراءات المنهجية:", bold=True)
        self._add_paragraph(
            f'"تم استخدام اختبار مربع كاي (Chi-Square Test) للكشف عن العلاقة بين المتغيرين الاسميين، '
            f'حيث بلغت العينة الكلية N = {results.get("N", "X")}. وقد تم اعتماد مستوى دلالة α = 0.05 '
            f'كمعيار للحكم على الدلالة الإحصائية."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        if results.get('دال'):
            self._add_paragraph(
                '"أظهرت نتائج اختبار مربع كاي وجود علاقة دالة إحصائياً بين [المتغير الأول] و[المتغير الثاني] '
                '(χ² = X.XX, p < 0.05), مما يدل على عدم استقلالية المتغيرين ووجود ارتباط بينهما."'
            )
        else:
            self._add_paragraph(
                '"أظهرت نتائج اختبار مربع كاي عدم وجود علاقة دالة إحصائياً بين [المتغير الأول] و[المتغير الثاني] '
                '(χ² = X.XX, p > 0.05), مما يدل على استقلالية المتغيرين."'
            )
        
        return self.doc
    
    def generate_regression(self, results):
        """Generate Multiple Linear Regression report - Enhanced"""
        self._add_title("تحليل الانحدار الخطي المتعدد\nMultiple Linear Regression")
        self.doc.add_paragraph()
        
        if 'error' in results:
            self._add_paragraph(f"❌ خطأ: {results['error']}")
            return self.doc
        
        # المعادلة الرياضية أولاً
        self._add_section_header("📐 المعادلة الرياضية للنموذج:")
        self._add_paragraph(
            "تمثل المعادلة التالية النموذج الرياضي للانحدار المتعدد المُستخرج من البيانات، "
            "حيث Y هو المتغير التابع، والمتغيرات X هي المتغيرات المستقلة، و ε هو حد الخطأ العشوائي."
        )
        self.doc.add_paragraph()
        
        # بناء المعادلة
        equation_parts = []
        constant = results.get('المعامل_الثابت', 0)
        equation_parts.append(f"Y = {constant:.3f}")
        
        for coef in results.get('معاملات', []):
            if coef['المتغير'] != 'الثابت':
                b_value = coef['المعامل']
                var_name = coef['المتغير']
                sign = "+" if b_value >= 0 else ""
                equation_parts.append(f" {sign} {b_value:.3f}({var_name})")
        
        equation = "".join(equation_parts) + " + ε"
        
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(equation)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        
        self.doc.add_paragraph()
        
        # ملخص النموذج
        self._add_section_header("📊 أولاً: ملخص النموذج - Model Summary")
        self._add_paragraph(
            "يوضح الجدول التالي جودة النموذج الإحصائي، حيث يُظهر معامل الارتباط المتعدد (R) "
            "ومعامل التحديد (R²) والخطأ المعياري للتقدير. معامل التحديد يوضح نسبة التباين "
            "في المتغير التابع التي يمكن تفسيرها بواسطة المتغيرات المستقلة."
        )
        self.doc.add_paragraph()
        
        table = self._create_table(rows=2, cols=4, headers=['R', 'R²', 'Adjusted R²', 'Std. Error'])
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], f"{results['R']:.3f}")
        self._fill_table_cell(cells[1], f"{results['R2']:.3f}")
        self._fill_table_cell(cells[2], f"{results['R2_المعدل']:.3f}")
        self._fill_table_cell(cells[3], f"{results['الخطأ_المعياري']:.3f}")
        
        self.doc.add_paragraph()
        
        # ANOVA للنموذج
        self._add_section_header("📈 ثانياً: اختبار معنوية النموذج - ANOVA")
        self._add_paragraph(
            "يختبر هذا الجدول ما إذا كان النموذج ككل دالاً إحصائياً أم لا، أي هل المتغيرات "
            "المستقلة مجتمعة لها تأثير دال على المتغير التابع."
        )
        self.doc.add_paragraph()
        
        table = self._create_table(rows=2, cols=3, headers=['F', 'df', 'Sig.'])
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], f"{results['F']:.3f}")
        self._fill_table_cell(cells[1], results.get('df', '-'))
        self._fill_table_cell(cells[2], f"{results['p_model']:.4f}")
        
        self.doc.add_paragraph()
        
        # معاملات الانحدار
        self._add_section_header("📋 ثالثاً: معاملات الانحدار - Coefficients")
        self._add_paragraph(
            "يعرض الجدول التالي معاملات الانحدار لكل متغير مستقل، حيث B هو المعامل غير المعياري، "
            "و t هو قيمة الاختبار، و Sig. هو مستوى الدلالة. تُظهر هذه القيم تأثير كل متغير مستقل "
            "على المتغير التابع بشكل منفرد."
        )
        self.doc.add_paragraph()
        
        num_vars = len(results.get('معاملات', []))
        table = self._create_table(rows=num_vars + 1, cols=4, headers=['المتغير', 'B', 't', 'Sig.'])
        
        for i, coef in enumerate(results.get('معاملات', []), start=1):
            cells = table.rows[i].cells
            self._fill_table_cell(cells[0], coef['المتغير'], align='right', bold=True)
            self._fill_table_cell(cells[1], f"{coef['المعامل']:.3f}")
            self._fill_table_cell(cells[2], f"{coef.get('t', 'N/A'):.3f}" if isinstance(coef.get('t'), (int, float)) else 'N/A')
            self._fill_table_cell(cells[3], f"{coef['p']:.4f}")
        
        self.doc.add_paragraph()
        
        # التفسير الأكاديمي المطول
        self._add_section_header("📖 رابعاً: التفسير الأكاديمي المفصل")
        
        if results.get('دال'):
            r2_percent = results['R2'] * 100
            
            interp = (
                f"أظهرت نتائج تحليل الانحدار الخطي المتعدد أن النموذج ككل دال إحصائياً عند مستوى دلالة 0.05, "
                f"حيث بلغت قيمة F المحسوبة ({results['F']:.3f}) بقيمة احتمالية (p = {results['p_model']:.4f}). "
                f"وهذا يعني أن المتغيرات المستقلة المُدرجة في النموذج لها تأثير دال إحصائياً على المتغير التابع.\n\n"
                f"كما بلغ معامل التحديد (R² = {results['R2']:.3f}), مما يشير إلى أن المتغيرات المستقلة "
                f"تفسر ما نسبته ({r2_percent:.1f}%) من التباين الكلي في المتغير التابع. "
                f"وهذه نسبة تعتبر {'جيدة' if results['R2'] >= 0.5 else 'مقبولة' if results['R2'] >= 0.3 else 'ضعيفة'} "
                f"في مجال العلوم الاجتماعية والإنسانية، حيث تتأثر الظواهر بعوامل متعددة ومعقدة.\n\n"
            )
            
            # تفسير المعاملات الدالة
            dalah_coefs = [c for c in results.get('معاملات', []) if c['p'] < 0.05 and c['المتغير'] != 'الثابت']
            
            if dalah_coefs:
                interp += "أما على مستوى المتغيرات المستقلة الفردية، فقد أظهرت النتائج ما يلي:\n\n"
                
                for coef in dalah_coefs:
                    direction = "إيجابي (طردي)" if coef['المعامل'] > 0 else "سلبي (عكسي)"
                    interp += (
                        f"• المتغير ({coef['المتغير']}): له تأثير {direction} دال إحصائياً على المتغير التابع "
                        f"(B = {coef['المعامل']:.3f}, t = {coef.get('t', 'N/A'):.3f}, p = {coef['p']:.4f}). "
                        f"وهذا يعني أن كل زيادة بمقدار وحدة واحدة في هذا المتغير تؤدي إلى "
                        f"{'زيادة' if coef['المعامل'] > 0 else 'نقصان'} في المتغير التابع بمقدار "
                        f"({abs(coef['المعامل']):.3f}) وحدة، مع ثبات العوامل الأخرى.\n\n"
                    )
            
            interp += (
                "\n\nمن الناحية العملية، يمكن استخدام هذا النموذج للتنبؤ بقيم المتغير التابع بناءً على "
                "قيم المتغيرات المستقلة. كما تساعد هذه النتائج في فهم الأهمية النسبية لكل متغير مستقل "
                "في التأثير على المتغير التابع، مما يوفر أساساً لاتخاذ القرارات أو بناء التوصيات."
            )
        else:
            interp = (
                f"أظهرت نتائج تحليل الانحدار الخطي المتعدد أن النموذج ككل غير دال إحصائياً عند مستوى دلالة 0.05, "
                f"حيث بلغت قيمة F المحسوبة ({results['F']:.3f}) بقيمة احتمالية (p = {results['p_model']:.4f}). "
                f"وهذا يعني أن المتغيرات المستقلة المُدرجة في النموذج ليس لها تأثير دال إحصائياً على المتغير التابع.\n\n"
                "من الناحية العملية، قد يشير هذا إلى أن المتغيرات المستقلة المختارة لا تفسر التباين "
                "في المتغير التابع بشكل كافٍ، أو أن حجم العينة غير كافٍ، أو أن العلاقة بين المتغيرات "
                "ليست خطية. وقد يتطلب الأمر إعادة النظر في اختيار المتغيرات أو استخدام نماذج أخرى."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        self._add_section_header("📝 خامساً: كيفية الكتابة في المذكرة")
        
        self._add_paragraph("• في فصل الإجراءات المنهجية:", bold=True)
        self._add_paragraph(
            '"تم استخدام تحليل الانحدار الخطي المتعدد (Multiple Linear Regression) لتحديد تأثير المتغيرات '
            'المستقلة على المتغير التابع. وقد تم اعتماد مستوى دلالة α = 0.05 كمعيار للحكم على دلالة النموذج '
            'والمعاملات الفردية."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        if results.get('دال'):
            self._add_paragraph(
                '"أظهرت نتائج تحليل الانحدار المتعدد أن النموذج دال إحصائياً (F = X.XX, p < 0.05), '
                'حيث فسّرت المتغيرات المستقلة ما نسبته (R² = X.XX) من التباين في المتغير التابع. '
                'كما أظهرت النتائج أن المتغير [اسم المتغير] له تأثير دال (B = X.XX, p < 0.05)."'
            )
        else:
            self._add_paragraph(
                '"أظهرت نتائج تحليل الانحدار المتعدد أن النموذج غير دال إحصائياً (F = X.XX, p > 0.05), '
                'مما يشير إلى أن المتغيرات المستقلة لا تفسر التباين في المتغير التابع بشكل دال."'
            )
        
        return self.doc
    
    def generate_ttest(self, results):
        """Generate T-Test report"""
        if 'error' in results:
            self._add_paragraph(f"خطأ: {results['error']}", color='red')
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 أولاً: معلومات التحليل")
        self._add_paragraph(f"• نوع الاختبار: اختبار T للعينات المستقلة (Independent Samples T-Test)")
        self._add_paragraph(f"• حجم العينة الكلي: N = {results['المجموعة_1']['العدد'] + results['المجموعة_2']['العدد']}")
        self._add_paragraph(f"• مستوى الدلالة المعتمد: α = 0.05")
        self.doc.add_paragraph()
        
        # الإحصاءات الوصفية
        self._add_section_header("📊 ثانياً: الإحصاءات الوصفية للمجموعات")
        
        table = self._create_table(
            rows=3,
            cols=4,
            headers=['المجموعة', 'N', 'Mean', 'Std. Deviation']
        )
        
        # المجموعة 1
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], results['المجموعة_1']['الاسم'], align='right', bold=True)
        self._fill_table_cell(cells[1], str(results['المجموعة_1']['العدد']))
        self._fill_table_cell(cells[2], f"{results['المجموعة_1']['المتوسط']:.2f}")
        self._fill_table_cell(cells[3], f"{results['المجموعة_1']['الانحراف']:.2f}")
        
        # المجموعة 2
        cells = table.rows[2].cells
        self._fill_table_cell(cells[0], results['المجموعة_2']['الاسم'], align='right', bold=True)
        self._fill_table_cell(cells[1], str(results['المجموعة_2']['العدد']))
        self._fill_table_cell(cells[2], f"{results['المجموعة_2']['المتوسط']:.2f}")
        self._fill_table_cell(cells[3], f"{results['المجموعة_2']['الانحراف']:.2f}")
        
        self.doc.add_paragraph()
        
        # نتائج اختبار T
        self._add_section_header("📊 ثالثاً: نتائج اختبار T")
        
        table = self._create_table(
            rows=2,
            cols=4,
            headers=['t', 'df', 'Sig. (2-tailed)', "Cohen's d"]
        )
        
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], f"{results['t']:.3f}")
        self._fill_table_cell(cells[1], str(results['df']))
        self._fill_table_cell(cells[2], f"{results['p']:.4f}")
        self._fill_table_cell(cells[3], f"{results['cohens_d']:.3f}")
        
        self.doc.add_paragraph()
        
        # التفسير الأكاديمي
        self._add_section_header("📖 رابعاً: التفسير الأكاديمي المفصل")
        
        if results['دال']:
            interp = (
                f"أظهرت نتائج اختبار T للعينات المستقلة وجود فروق ذات دلالة إحصائية بين المجموعتين "
                f"عند مستوى دلالة 0.05, حيث بلغت قيمة t المحسوبة ({results['t']:.3f}) بدرجات حرية "
                f"({results['df']}), وبقيمة احتمالية p = {results['p']:.4f}.\n\n"
                
                f"كما بلغ حجم الأثر (Cohen's d = {results['cohens_d']:.3f}) وهو يُصنف على أنه "
                f"{results['حجم_الأثر']}, مما يشير إلى أن الفرق بين المجموعتين {results['حجم_الأثر']} من الناحية العملية.\n\n"
                
                f"من الناحية العملية، تشير هذه النتائج إلى وجود اختلاف حقيقي وملموس بين المجموعتين، "
                f"حيث كان متوسط المجموعة الأولى ({results['المجموعة_1']['المتوسط']:.2f}) "
                f"{'أعلى' if results['المجموعة_1']['المتوسط'] > results['المجموعة_2']['المتوسط'] else 'أقل'} "
                f"من متوسط المجموعة الثانية ({results['المجموعة_2']['المتوسط']:.2f})."
            )
        else:
            interp = (
                f"أظهرت نتائج اختبار T للعينات المستقلة عدم وجود فروق ذات دلالة إحصائية بين المجموعتين "
                f"عند مستوى دلالة 0.05, حيث بلغت قيمة t المحسوبة ({results['t']:.3f}) بدرجات حرية "
                f"({results['df']}), وبقيمة احتمالية p = {results['p']:.4f}.\n\n"
                
                f"وهذا يعني أن الفرق الظاهري بين متوسط المجموعة الأولى ({results['المجموعة_1']['المتوسط']:.2f}) "
                f"والمجموعة الثانية ({results['المجموعة_2']['المتوسط']:.2f}) ليس دالاً إحصائياً، "
                f"وقد يكون ناتجاً عن الصدفة أو التباين العشوائي في العينة."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        self._add_section_header("✍️ خامساً: دليل الكتابة في المذكرة")
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل المنهجية:", bold=True)
        self._add_paragraph(
            '"للإجابة على [السؤال/الفرضية]، تم استخدام اختبار T للعينات المستقلة (Independent Samples T-Test) '
            'لمقارنة المتوسطات بين مجموعتين مستقلتين. تم اعتماد مستوى دلالة α = 0.05 للحكم على الدلالة الإحصائية '
            'للفروق بين المجموعات."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        if results['دال']:
            self._add_paragraph(
                f'"أظهرت نتائج اختبار T وجود فروق دالة إحصائياً بين المجموعتين (t = {results["t"]:.3f}, '
                f'df = {results["df"]}, p = {results["p"]:.4f}), حيث كان متوسط [المجموعة الأولى] '
                f'أعلى/أقل من متوسط [المجموعة الثانية] بفارق دال إحصائياً."'
            )
        else:
            self._add_paragraph(
                f'"أظهرت نتائج اختبار T عدم وجود فروق دالة إحصائياً بين المجموعتين (t = {results["t"]:.3f}, '
                f'df = {results["df"]}, p = {results["p"]:.4f}), مما يشير إلى تشابه المجموعتين في المتغير المدروس."'
            )
        
        return self.doc
    
    def generate_cronbach(self, results):
        """Generate Cronbach's Alpha report"""
        if 'error' in results:
            self._add_paragraph(f"خطأ: {results['error']}", color='red')
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 أولاً: معلومات التحليل")
        self._add_paragraph(f"• نوع الاختبار: معامل ألفا كرونباخ (Cronbach's Alpha)")
        self._add_paragraph(f"• عدد البنود (Items): N = {results['عدد_البنود']}")
        self._add_paragraph(f"• حجم العينة: N = {results.get('N', 'غير محدد')}")
        self.doc.add_paragraph()
        
        # نتيجة ألفا
        self._add_section_header("📊 ثانياً: نتيجة معامل ألفا كرونباخ")
        
        table = self._create_table(
            rows=2,
            cols=2,
            headers=["Cronbach's Alpha", 'N of Items']
        )
        
        cells = table.rows[1].cells
        self._fill_table_cell(cells[0], f"{results['alpha']:.3f}")
        self._fill_table_cell(cells[1], str(results['عدد_البنود']))
        
        self.doc.add_paragraph()
        
        # التفسير الأكاديمي
        self._add_section_header("📖 ثالثاً: التفسير الأكاديمي المفصل")
        
        alpha_val = results['alpha']
        if alpha_val >= 0.9:
            quality = "ممتازة جداً"
        elif alpha_val >= 0.8:
            quality = "جيدة"
        elif alpha_val >= 0.7:
            quality = "مقبولة"
        elif alpha_val >= 0.6:
            quality = "مقبولة بشكل حدي"
        else:
            quality = "ضعيفة"
        
        interp = (
            f"بلغت قيمة معامل ألفا كرونباخ (α = {results['alpha']:.3f}) للمقياس المكون من "
            f"{results['عدد_البنود']} بنداً، وهي قيمة تُصنف على أنها {quality} حسب معايير "
            f"جورج ومالري (George & Mallery, 2003).\n\n"
            
            f"يشير ذلك إلى أن المقياس يتمتع بدرجة {quality} من الاتساق الداخلي، مما يعني أن "
            f"البنود المكونة للمقياس {'تقيس بشكل متسق نفس المفهوم' if alpha_val >= 0.7 else 'قد لا تقيس نفس المفهوم بشكل كافٍ'}. "
        )
        
        if alpha_val >= 0.7:
            interp += (
                f"وهذا يدعم استخدام المقياس في الدراسة الحالية كأداة موثوقة لقياس المتغير المستهدف."
            )
        else:
            interp += (
                f"وقد يستدعي ذلك مراجعة بنود المقياس أو حذف بعض البنود التي قد تقلل من الاتساق الداخلي."
            )
        
        self._add_paragraph(interp)
        
        # دليل الكتابة
        self.doc.add_paragraph()
        self._add_section_header("✍️ رابعاً: دليل الكتابة في المذكرة")
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل المنهجية:", bold=True)
        self._add_paragraph(
            '"للتحقق من ثبات المقياس، تم حساب معامل ألفا كرونباخ (Cronbach\'s Alpha)، '
            'وهو مؤشر يقيس الاتساق الداخلي للمقياس، ويتراوح بين 0 و 1. '
            'القيم الأعلى من 0.7 تُعتبر مقبولة أكاديمياً."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        if alpha_val >= 0.7:
            self._add_paragraph(
                f'"أظهرت النتائج أن المقياس يتمتع بثبات {quality} (α = {results["alpha"]:.3f})، '
                f'مما يدعم استخدامه في الدراسة الحالية."'
            )
        else:
            self._add_paragraph(
                f'"أظهرت النتائج أن المقياس يتمتع بثبات {quality} (α = {results["alpha"]:.3f})، '
                f'مما قد يستدعي مراجعة بنوده أو تحسينه في الدراسات المستقبلية."'
            )
        
        return self.doc
    
    def generate_descriptive(self, results):
        """Generate Descriptive Statistics report"""
        if 'error' in results:
            self._add_paragraph(f"خطأ: {results['error']}", color='red')
            return self.doc
        
        # معلومات التحليل
        self._add_section_header("📋 أولاً: معلومات التحليل")
        self._add_paragraph("• نوع التحليل: الإحصاء الوصفي (Descriptive Statistics)")
        
        total_vars = 0
        if 'متغيرات_رقمية' in results:
            total_vars += len(results['متغيرات_رقمية'])
        if 'متغيرات_فئوية' in results:
            total_vars += len(results['متغيرات_فئوية'])
        
        self._add_paragraph(f"• عدد المتغيرات المدروسة: {total_vars}")
        self.doc.add_paragraph()
        
        # المتغيرات الرقمية
        if 'متغيرات_رقمية' in results and results['متغيرات_رقمية']:
            self._add_section_header("📊 ثانياً: الإحصاءات الوصفية للمتغيرات الرقمية")
            
            table = self._create_table(
                rows=len(results['متغيرات_رقمية']) + 1,
                cols=6,
                headers=['المتغير', 'N', 'Mean', 'Std. Deviation', 'Min', 'Max']
            )
            
            for i, var in enumerate(results['متغيرات_رقمية'], start=1):
                cells = table.rows[i].cells
                self._fill_table_cell(cells[0], var['المتغير'], align='right', bold=True)
                self._fill_table_cell(cells[1], str(var['العدد']))
                self._fill_table_cell(cells[2], f"{var['المتوسط']:.2f}")
                self._fill_table_cell(cells[3], f"{var['الانحراف_المعياري']:.2f}")
                self._fill_table_cell(cells[4], f"{var['أصغر_قيمة']:.2f}")
                self._fill_table_cell(cells[5], f"{var['أكبر_قيمة']:.2f}")
            
            self.doc.add_paragraph()
            
            # تفسير مختصر
            self._add_paragraph(
                "يعرض الجدول أعلاه ملخصاً للإحصاءات الوصفية للمتغيرات الرقمية، حيث يتضمن "
                "حجم العينة (N)، المتوسط الحسابي (Mean)، الانحراف المعياري (Std. Deviation)، "
                "أصغر قيمة (Min)، وأكبر قيمة (Max) لكل متغير."
            )
        
        # المتغيرات الفئوية
        if 'متغيرات_فئوية' in results and results['متغيرات_فئوية']:
            self.doc.add_paragraph()
            section_num = "ثالثاً" if 'متغيرات_رقمية' in results else "ثانياً"
            self._add_section_header(f"📊 {section_num}: التوزيعات التكرارية للمتغيرات الفئوية")
            
            for var_data in results['متغيرات_فئوية']:
                self.doc.add_paragraph()
                self._add_paragraph(f"• {var_data['المتغير']}:", bold=True)
                
                table = self._create_table(
                    rows=len(var_data['توزيع']) + 1,
                    cols=3,
                    headers=['الفئة', 'Frequency', 'Percent']
                )
                
                for i, (category, freq, percent) in enumerate(var_data['توزيع'], start=1):
                    cells = table.rows[i].cells
                    self._fill_table_cell(cells[0], str(category), align='right')
                    self._fill_table_cell(cells[1], str(freq))
                    self._fill_table_cell(cells[2], f"{percent:.1f}%")
                
                self.doc.add_paragraph()
        
        # دليل الكتابة
        self.doc.add_paragraph()
        next_section = "رابعاً" if ('متغيرات_رقمية' in results and 'متغيرات_فئوية' in results) else "ثالثاً"
        self._add_section_header(f"✍️ {next_section}: دليل الكتابة في المذكرة")
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل المنهجية:", bold=True)
        self._add_paragraph(
            '"تم استخدام الإحصاء الوصفي (Descriptive Statistics) لوصف خصائص العينة '
            'والمتغيرات المدروسة، حيث تم حساب المتوسطات الحسابية والانحرافات المعيارية '
            'للمتغيرات الرقمية، والتوزيعات التكرارية والنسب المئوية للمتغيرات الفئوية."'
        )
        
        self.doc.add_paragraph()
        self._add_paragraph("• في فصل النتائج:", bold=True)
        self._add_paragraph(
            '"أظهرت نتائج الإحصاء الوصفي أن [وصف مختصر للنتائج الرئيسية، مثل متوسطات المتغيرات '
            'أو التوزيعات الأكثر شيوعاً]."'
        )
        
        return self.doc
    
    def save(self, filename):
        """Save document to file"""
        self.doc.save(filename)
        return filename
